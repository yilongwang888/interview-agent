from pypdf import PdfReader
from langchain.prompts.chat import SystemMessagePromptTemplate,HumanMessagePromptTemplate,ChatPromptTemplate,MessagesPlaceholder
from operator import itemgetter

def resume_evaluation_bot(text,llm):
    '''
    @function:对求职者的简历文本进行打分与评估
    
    @parmas:
    text:str 求职者的简历文本
    llm :langchain_openai.chat_models.base.ChatOpenAI 大语言模型

    @return:AIMessage 输出评估结果字符串
    '''
    
    #prompt工程
    system_prompt = SystemMessagePromptTemplate.from_template("你是一个简历评估专家,合理公正地为求职者的简历进行打分和评估")
    user_prompt   = HumanMessagePromptTemplate.from_template("""评估的时候基于以下简历文本内容：{context} 问题：{query} """)
    full_chat_prompt = ChatPromptTemplate.from_messages([system_prompt,MessagesPlaceholder(variable_name='chat_history'),user_prompt])

    #chat链
    chat_chain = {
        "context": itemgetter("context"),
        "query"  : itemgetter("query"),
        "chat_history":itemgetter("chat_history")
         } | full_chat_prompt | llm

    query = "根据我的简历，客观以技术能力，问题解决能力，项目管理，软技能，工作经验给我评分（满分100），注意要先输出分数列表再说明原因，\
             即先给出5个分数再分别说明原因。输出举例[90,80,100,85,92]，换行，解决能力分数是90原因是XXXXX，项目管理分数是80原因是XXX"
    response = chat_chain.invoke({'context':text,'query':query,'chat_history':[]})

    return response
    
def get_resume_evaluation(pdf_path:str,llm):
    '''
    @function:获取求职者简历pdf的能力评估

    @params:
    pdf_path: str 简历pdf的路径
    llm :langchain_openai.chat_models.base.ChatOpenAI 大语言模型

    @return: turple (list,list) 左为用户简历五方位评分，右为用户简历评估内容
    '''
    #获取简历pdf文本
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text().replace('\n',' ')

    #进行简历评估
    evaluation = resume_evaluation_bot(text,llm)
    score_list = evaluation.content.split('\n')[0]
    evalu_list = evaluation.content.split('\n')[1:]
    evalu_list = [e for e in evalu_list if e != '']
    return eval(score_list),evalu_list

import numpy as np
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon, Circle
from matplotlib.patheffects import withStroke, SimpleLineShadow


def resume_evaluation(values):

    labels = np.array(['技术能力', '项目经验', '学历背景', '软技能', '沟通能力'])
    values = np.concatenate((values, [values[0]]))      
    angles = np.linspace(0, 2*np.pi, len(labels), endpoint=False)
    angles = np.concatenate((angles, [angles[0]]))

    mpl.rcParams['font.family'] = 'WenQuanYi Micro Hei'               
    #mpl.rcParams['font.family'] = 'Noto Sans CJK SC'
    mpl.rcParams['axes.unicode_minus'] = False
    fig = plt.figure(figsize=(6, 6), dpi=120)           
    ax = fig.add_subplot(111, polar=True)
    ax.set_theta_zero_location('N')
    ax.set_ylim(0, 100)


    ax.set_thetagrids(angles[:-1]*180/np.pi, labels, fontsize=8,  
                  color='#B4C8EA')                     
    ax.set_rticks([20, 40, 60, 80, 100])
    ax.tick_params(axis='y', labelsize=6, colors='silver')  
    
    ax.set_rgrids([0, 20, 40, 60, 80, 100], ['']*6, alpha=0.25)
    ax.grid(axis='y', ls='-', lw=0.3, color='silver', alpha=0.25)
    ax.grid(axis='x', ls='-', lw=0.3, color='silver', alpha=0.25)


    bg_circle = Circle((0, 0), 100, transform=ax.transData,
                   fill=True, color='white', zorder=0)
    ax.add_patch(bg_circle)


    for alpha, scale in zip(np.linspace(0.55, 0.05, 15),
                        np.linspace(1, 0.2, 15)):
        xs, ys = angles, values * scale
        poly = Polygon(np.c_[xs, ys], closed=True,
                   fill=True, facecolor='#FF6B6B',
                   alpha=alpha, lw=0, zorder=2)
        ax.add_patch(poly)


    # 5.1 外层发光
    glow = ax.plot(angles, values, color='#FF6B6B', lw=2, zorder=3)[0]  
    glow.set_path_effects([withStroke(foreground='white', linewidth=4)])
    # 5.2 内层实线
    line = ax.plot(angles, values, color='#FF3B3B', lw=1, zorder=4)[0]  


    for a, v in zip(angles, values):
        ax.scatter(a, v, s=40, color='#FF3B3B', zorder=5, 
               edgecolors='white', linewidths=0.8)  
        # 内发光小圈
        ax.scatter(a, v, s=80, color='#FF3B3B', zorder=4,
               alpha=0.15, linewidth=0)


    # 7.1 渐变色映射
    cmap = mpl.colormaps['cool']
    title = ax.set_title('个人技能评估', size=12, pad=15,  
                     fontweight='bold')
    # 将标题字符逐字染色 → 伪渐变
    title.set_color(cmap(0.6))
    title.set_path_effects([SimpleLineShadow(offset=(1, -1),
                                         alpha=0.3, linewidth=0),
                        withStroke(foreground='white', linewidth=1)])  

    # 7.2 动态副标题
    ax.text(0.5, 1.1, f'综合得分 {int(values.mean())} 分',
        transform=ax.transAxes, ha='center', va='center',
        fontsize=6, color='grey')  # 副标题字体调小


    plt.tight_layout()
    plt.savefig('radar_chart.png', dpi=120, bbox_inches='tight')  
    #plt.show()

import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Mm
import os
import subprocess
from datetime import date

def add_text_watermark(doc, watermark_text="YiLong AI"):
    """添加斜向文字水印（修复XML解析错误）"""
    try:
        section = doc.sections[0]
        sectPr = section._sectPr
        
        # 清理已有水印
        for watermark in sectPr.findall('.//w:watermark', namespaces=sectPr.nsmap):
            sectPr.remove(watermark)
        
        # 修复XML格式：单行紧凑写法，避免解析错误
        watermark_xml = (
            '<w:watermark xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:textWatermark w:font="微软雅黑" w:color="D9D9D9" w:semitransparent="1" '
            f'w:scale="200" w:layout="diagonal" w:text="{watermark_text}"/>'
            '</w:watermark>'
        )
        
        # 解析XML并添加
        watermark = parse_xml(watermark_xml)
        sectPr.append(watermark)
        print("✅ 斜向水印「YiLong AI」添加成功")
    except Exception as e:
        print(f"⚠️ 水印添加失败: {e}")
        # 降级方案：使用英文水印（避免中文XML解析问题）
        try:
            fallback_xml = (
                '<w:watermark xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:textWatermark w:font="Arial" w:color="D9D9D9" w:semitransparent="1" '
                'w:scale="200" w:layout="diagonal" w:text="YiLong AI"/>'
                '</w:watermark>'
            )
            watermark = parse_xml(fallback_xml)
            sectPr.append(watermark)
            print("✅ 降级水印「YiLong AI」添加成功（Arial字体）")
        except:
            print("❌ 所有水印添加方式均失败")

def add_header_footer(doc):
    """添加简洁的页脚（版权信息）"""
    # 清空页眉（保持干净）
    header = doc.sections[0].header
    for para in header.paragraphs:
        para.clear()
    
    # 页脚：居中版权信息
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('© 2026 YiLong AI 版权所有')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(153, 153, 153)

def resume_assessment_report(body_text: list, score_list: list, img_path: str):
    """
    稳定版：AI简历测评报告（YiLong AI斜向文字水印）
    """
    # 安全参数处理
    body_text = body_text if isinstance(body_text, list) else []
    score_list = [float(x) for x in score_list if isinstance(x, (int, float))]
    avg_score = round(sum(score_list)/len(score_list), 1) if score_list else 0.0

    # ======================
    # 1. 初始化文档 + 全局设置
    # ======================
    doc = docx.Document()
    
    # A4纸张标准设置
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 全局字体：微软雅黑（兼容中文）
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Microsoft YaHei' if os.name == 'nt' else 'SimSun'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑' if os.name == 'nt' else '宋体')
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)

    # ======================
    # 2. 添加水印 + 页眉页脚
    # ======================
    add_text_watermark(doc, "YiLong AI")  # 核心：斜向YiLong AI水印
    add_header_footer(doc)

    # ======================
    # 3. 标题区域（大气简洁）
    # ======================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('AI简历智能测评报告')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
    title_para.paragraph_format.space_after = Pt(15)

    # 分割线
    line_para = doc.add_paragraph("="*50)
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    line_para.paragraph_format.space_after = Pt(20)

    # ======================
    # 4. 评分展示（视觉突出）
    # ======================
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_para.add_run(f'综合评分：{avg_score} 分')
    score_run.font.size = Pt(18)
    score_run.font.bold = True
    # 分数颜色区分
    if avg_score >= 80:
        score_run.font.color.rgb = RGBColor(0, 128, 0)    # 绿色-优秀
    elif avg_score >= 60:
        score_run.font.color.rgb = RGBColor(255, 140, 0) # 橙色-合格
    else:
        score_run.font.color.rgb = RGBColor(204, 0, 0)   # 红色-不合格
    score_para.paragraph_format.space_after = Pt(15)

    # ======================
    # 5. 测评结果提示
    # ======================
    result_para = doc.add_paragraph()
    result_para.paragraph_format.first_line_indent = Inches(0.3)
    result_para.paragraph_format.line_spacing = 1.5
    
    if avg_score > 80:
        result_run = result_para.add_run('🎉 恭喜您！您已通过本次AI简历测评，综合表现优秀。')
        result_run.font.color.rgb = RGBColor(0, 102, 0)
    else:
        result_run = result_para.add_run('💡 很遗憾，您暂未达到优选标准，可根据报告建议优化简历后重新投递。')
        result_run.font.color.rgb = RGBColor(204, 0, 0)
    result_run.font.bold = True
    result_para.paragraph_format.space_after = Pt(15)

    # ======================
    # 6. 详细评价
    # ======================
    if body_text:
        detail_title = doc.add_paragraph('📋 详细评价')
        detail_title.runs[0].font.size = Pt(14)
        detail_title.runs[0].font.bold = True
        detail_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        detail_title.paragraph_format.space_after = Pt(10)

        for idx, text in enumerate(body_text):
            if text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.3)
                p.paragraph_format.line_spacing = 1.5
                p.add_run(f'{idx+1}. ').font.bold = True
                p.add_run(text.strip())

    doc.add_paragraph()

    # ======================
    # 7. 评分可视化图表
    # ======================
    chart_title = doc.add_paragraph('📊 评分维度可视化')
    chart_title.runs[0].font.size = Pt(14)
    chart_title.runs[0].font.bold = True
    chart_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    chart_title.paragraph_format.space_after = Pt(10)

    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        if os.path.exists(img_path) and os.path.getsize(img_path) > 100:
            img_para.add_run().add_picture(img_path, width=Inches(6.0))
        else:
            img_para.add_run('[ 评分图表暂未生成 ]').font.color.rgb = RGBColor(204, 0, 0)
    except:
        img_para.add_run('[ 图表加载失败 ]').font.color.rgb = RGBColor(204, 0, 0)
    img_para.paragraph_format.space_after = Pt(15)

    # ======================
    # 8. 后续指引
    # ======================
    guide_title = doc.add_paragraph('🔍 下一步指引')
    guide_title.runs[0].font.size = Pt(14)
    guide_title.runs[0].font.bold = True
    guide_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    guide_title.paragraph_format.space_after = Pt(10)

    guide_para = doc.add_paragraph()
    guide_para.paragraph_format.first_line_indent = Inches(0.3)
    guide_para.paragraph_format.line_spacing = 1.5

    if avg_score > 80:
        guide_para.add_run('请于10月1日24:00前访问以下链接完成AI智能面试：\n')
        url_run = guide_para.add_run('http://yilongweiwang.com')
        url_run.font.color.rgb = RGBColor(0, 51, 255)
        url_run.font.underline = True
        guide_para.add_run('\n祝您面试顺利，期待您的加入！')
    else:
        guide_para.add_run('建议您根据上述评价优化简历内容，重点完善项目成果量化、核心技能突出等方面，优化完成后可重新投递。\n下一次，期待看到更优秀的你！')

    # ======================
    # 9. 落款
    # ======================
    end_line = doc.add_paragraph("─"*60)
    end_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_line.paragraph_format.space_after = Pt(10)

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.add_run(f'{date.today().strftime("%Y年%m月%d日")}\n')
    footer_para.add_run('YiLong AI 项目组').font.bold = True
    footer_para.add_run(' 出品')

    # ======================
    # 10. 保存并转换PDF
    # ======================
    current_dir = os.getcwd()
    word_tmp = os.path.join(current_dir, 'temp_report.docx')
    pdf_out = os.path.join(current_dir, 'AI简历测评报告.pdf')

    try:
        # 保存Word文档
        doc.save(word_tmp)
        
        # 转换PDF（兼容多系统）
        libreoffice_paths = ['libreoffice', '/usr/bin/libreoffice', 
                           '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                           'C:/Program Files/LibreOffice/program/soffice.exe']
        convert_success = False
        
        for lo_path in libreoffice_paths:
            try:
                subprocess.run([lo_path, '--headless', '--convert-to', 'pdf',
                               '--outdir', current_dir, word_tmp],
                              check=True, capture_output=True, timeout=30)
                convert_success = True
                break
            except:
                continue
        
        if convert_success:
            print(f"✅ PDF报告生成成功：{pdf_out}")
        else:
            print(f"⚠️ PDF转换失败，Word文件已保存：{word_tmp}")
            
    except Exception as e:
        print(f"❌ 保存失败：{e}")
    finally:
        # 清理临时文件
        try:
            if os.path.exists(word_tmp):
                os.remove(word_tmp)
        except:
            pass
